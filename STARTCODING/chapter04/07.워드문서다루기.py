from docx import Document

# 1. 워드 생성
document = Document()

# 2. 워드 데이터 추가
document.add_heading("기사 제목", level = 0)
document.add_paragraph("기사 링크")
document.add_paragraph("기사 본문")

# 3. 워드 저장
document.save("test.docx")